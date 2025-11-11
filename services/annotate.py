from docx import Document
from docx.enum.text import WD_COLOR_INDEX

def _severity_to_highlight(sev: str):
    s = (sev or "").lower()
    if s == "high":
        return WD_COLOR_INDEX.YELLOW
    if s == "medium":
        return WD_COLOR_INDEX.BRIGHT_GREEN
    return WD_COLOR_INDEX.GRAY_25  # low/default

def build_annotated_docx(translation_segments, issues, out_path):
    """
    Create a DOCX:
      - Each target segment on its own paragraph
      - If the segment has issues, highlight the entire paragraph
      - Append inline note with issue shortcodes
      - Then add a bullet list with details (type + evidence [+ suggestion])
    """
    doc = Document()
    issue_map = {}
    for it in issues or []:
        seg = it.get("segment", 0)
        issue_map.setdefault(seg, []).append(it)

    for idx, tgt in enumerate(translation_segments, start=1):
        p = doc.add_paragraph()
        run = p.add_run(tgt if tgt else "")
        if idx in issue_map:
            sev_top = "low"
            types = []
            for it in issue_map[idx]:
                types.append(it.get("type", "issue"))
                sev = (it.get("severity") or "low").lower()
                if sev == "high":
                    sev_top = "high"
                elif sev == "medium" and sev_top != "high":
                    sev_top = "medium"
            run.font.highlight_color = _severity_to_highlight(sev_top)
            p.add_run(f"  [ISSUES: {', '.join(types)}]")

            for it in issue_map[idx]:
                detail = it.get("detail") or {}
                evidence = detail.get("evidence", "")
                suggestion = detail.get("suggestion", "")
                bullet = doc.add_paragraph(style="List Bullet")
                txt = f"Segment {idx} — {it.get('type','issue')} ({it.get('severity','low')}): "
                if evidence:
                    txt += f"{evidence}"
                if suggestion:
                    txt += f" | Suggestion: {suggestion}"
                bullet.add_run(txt)
    doc.save(out_path)

def save_plain_docx(translation_segments, out_path):
    """
    Save a clean, fixed translation DOCX (no highlights, just the edited segments).
    """
    doc = Document()
    for tgt in translation_segments:
        doc.add_paragraph(tgt or "")
    doc.save(out_path)
